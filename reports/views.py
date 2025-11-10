from django.shortcuts import render
from django.http import JsonResponse
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet,ParagraphStyle
from docx import Document
import google.generativeai as genai
import os
from django.conf import settings
import re
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
import calendar
from datetime import datetime

# Configuration Gemini
genai.configure(api_key=os.getenv("GEMINI_API_KEY", "AIzaSyCBLgdFqqf0DVB5zhGcNTfPc8p8vN1CX2w"))

def clean_content(text):
    paragraphs = text.split("\n\n")
    if len(paragraphs) > 2:
        paragraphs = paragraphs[2:]  # garder à partir du 3ème paragraphe
    text = "\n\n".join(paragraphs) 
    text = re.sub(r"\n{3,}", "\n\n", text) 
    text = re.sub(r"[*#_]", "", text) 
    text = re.sub(r"[::]", "", text)
    text = re.sub(r"[---]", "", text)
    return text.strip()

def parse_table(section):
    lines = [l.rstrip() for l in section.split("\n") if l.strip()]
    # Doit contenir au moins une ligne avec '|' (entête) et une autre ligne
    if len(lines) < 2 or "|" not in lines[0]:
        return None

    rows = []
    for l in lines:
        # Filtrer les lignes de séparation Markdown (ex: |---|---|---|)
        if re.match(r"^\s*\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?\s*$", l):
            continue
        # Si la ligne n'a pas de | mais est présente, on la considère comme simple texte
        if "|" not in l:
            continue
        parts = [p.strip() for p in l.split("|")]
        # Enlever éventuels éléments vides générés par | en début/fin
        # mais garder cellules vides intermédiaires si nécessaire
        if parts and any(cell != "" for cell in parts):
            # trim leading/trailing empty cells
            if parts[0] == "":
                parts = parts[1:]
            if parts and parts[-1] == "":
                parts = parts[:-1]
            rows.append(parts)
    if not rows:
        return None
    max_cols = max(len(r) for r in rows)
    normalized = []
    for r in rows:
        if len(r) < max_cols:
            r = r + [""] * (max_cols - len(r))
        normalized.append(r)
    return normalized

def add_page_number(canvas, doc, logo_path='C:/Users/dell/myproject/media/reports/logo.jpeg'):
    canvas.saveState()
    page_num = f"Page {canvas.getPageNumber()}"
    canvas.setFont('Helvetica', 9)
    canvas.drawRightString(A4[0] - 2*cm, 1.2*cm, page_num)
    if logo_path and os.path.exists(logo_path):
        try:
            canvas.drawImage(logo_path, 2*cm, A4[1] - 3*cm, width=3*cm, height=2*cm, preserveAspectRatio=True)
        except Exception:
            pass
    canvas.restoreState()

def get_temp_data():
    """données temporaire (janvier-février 2025)"""
    sites = [{"id": 1, "nom": "Casablanca Plant", "societe": "Innovatech Solutions", "region": "Grand Casablanca"}]

    energie = [
        {"site_id": 1, "annee": 2025, "mois": "Janvier",  "type": "Électricité", "unite": "kWh", "valeur": 15800},
        {"site_id": 1, "annee": 2025, "mois": "Février", "type": "Électricité", "unite": "kWh", "valeur": 15000},
    ]

    eau = [
        {"site_id": 1, "annee": 2025, "mois": "Janvier", "famille_culture": "Tomate", "variete": "Cerise", "eau_m3": 3500},
        {"site_id": 1, "annee": 2025, "mois": "Février", "famille_culture": "Tomate", "variete": "Cerise", "eau_m3": 3400},
    ]

    dechets = [
        {"site_id": 1, "annee": 2025, "mois": "Janvier", "categorie_dechets": "Plastique", "unite": "kg", "valeur": 480},
        {"site_id": 1, "annee": 2025, "mois": "Février", "categorie_dechets": "Métal", "unite": "kg", "valeur": 320},
    ]

    social = [
        {"site_id": 1, "annee": 2025, "mois": "Janvier", "action": "Formation sécurité", "budget": 2500, "nb": 40},
        {"site_id": 1, "annee": 2025, "mois": "Février", "action": "Don de sang", "budget": 1800, "nb": 60},
    ]

    production = [
        {"site_id": 1, "annee": 2025, "mois": "Janvier", "famille_culture": "Tomate", "variete": "Cerise", "sup_ha": 12, "prod_kg": 26500},
        {"site_id": 1, "annee": 2025, "mois": "Février", "famille_culture": "Tomate", "variete": "Cerise", "sup_ha": 12, "prod_kg": 28000},
    ]

    return {"sites": sites, "energie": energie, "eau": eau, "dechets": dechets, "social": social, "production": production}


def generate_report(request):
    if request.method == "POST":
        report_type = request.POST.get("report_type")
        referentiel = request.POST.get("referentiel")
        filters = request.POST.get("filters")
        start_str = request.POST.get("start")
        end_str = request.POST.get("end")

        try:
            start_date = datetime.strptime(start_str, "%Y-%m")
            end_year, end_month = map(int, end_str.split('-'))
            last_day = calendar.monthrange(end_year, end_month)[1]
            end_date = datetime(end_year, end_month, last_day)
        except Exception as e:
            return JsonResponse({"error": f"Format de date invalide: {str(e)}"}, status=400)

        try:
            data = get_temp_data()

            # --- Création du résumé pour le prompt ---
            summary_lines = ["Synthèse des données (janvier-février 2025) :"]
            for site in data["sites"]:
                sid = site["id"]
                sname = site["nom"]

                e_jan = sum(e["valeur"] for e in data["energie"] if e["site_id"]==sid and e["mois"]=="Janvier")
                e_fev = sum(e["valeur"] for e in data["energie"] if e["site_id"]==sid and e["mois"]=="Février")
                w_jan = sum(w["eau_m3"] for w in data["eau"] if w["site_id"]==sid and w["mois"]=="Janvier")
                w_fev = sum(w["eau_m3"] for w in data["eau"] if w["site_id"]==sid and w["mois"]=="Février")
                d_jan = sum(d["valeur"] for d in data["dechets"] if d["site_id"]==sid and d["mois"]=="Janvier")
                d_fev = sum(d["valeur"] for d in data["dechets"] if d["site_id"]==sid and d["mois"]=="Février")
                p_jan = sum(p["prod_kg"] for p in data["production"] if p["site_id"]==sid and p["mois"]=="Janvier")
                p_fev = sum(p["prod_kg"] for p in data["production"] if p["site_id"]==sid and p["mois"]=="Février")

                summary_lines.append(f"- {sname} : énergie (kWh) — Jan: {e_jan}, Fév: {e_fev}")
                summary_lines.append(f"  Production (kg) — Jan: {p_jan}, Fév: {p_fev}")
                summary_lines.append(f"  Eau (m³) — Jan: {w_jan}, Fév: {w_fev}")
                summary_lines.append(f"  Déchets (kg) — Jan: {d_jan}, Fév: {d_fev}")

            data_summary = "\n".join(summary_lines)

            # --- Génération avec Gemini ---
            model = genai.GenerativeModel("gemini-2.5-pro")
            if report_type == "extra-financier":
                prompt = f"""
                Tu es un rédacteur expert RSE. Rédige un rapport extra-financier conforme au référentiel {referentiel} 
                pour la période {start_date.strftime('%B %Y')} à {end_date.strftime('%B %Y')}. 
                Inclut les titres suivants :
                       - Résumé exécutif (100-150 mots) 
                       - Méthodologie (sources de données) 
                       - Section Environnement (KPIs, analyse) 
                       - Section Social (KPIs, actions) 
                       - Gouvernance 
                       - Alignement avec les Objectifs de Développement Durable (ODD)
                       - Recommandations
                       - Conclusion
                Format : sections avec titres.
                Données : {data_summary}
                """
            else:
                prompt = f"""
                Rédige un rapport de performance sous forme des tableaux comparatifs pour les indicateurs {filters}, 
                entre {start_date.strftime('%B %Y')} et {end_date.strftime('%B %Y')}.
                 Inclut les titres suivants :
                       -Résumé exécutif (100-150 mots) 
                       -Performance de Production et Consommation
                       -Performance RSE
                       -Conclusion 
                Données : {data_summary}
                """

            response = model.generate_content(prompt)
            # selon l'API, response peut contenir .text ou .output_text; on garde .text comme auparavant
            content = clean_content(getattr(response, "text", str(response)).strip())
            if not content:
                content = "Pas de contenu généré par Gemini."

        except Exception as e:
            return JsonResponse({"error": f"Erreur Gemini : {str(e)}"}, status=500)

        # --- Répertoire de sortie ---
        output_dir = os.path.join(settings.MEDIA_ROOT, "reports")
        os.makedirs(output_dir, exist_ok=True)
        base_name = f"{report_type}_{start_str}_{end_str}".replace(" ", "_")

        # --- Styles PDF ---
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(name="Title", fontSize=16, leading=18, spaceBefore=12, spaceAfter=12, textColor="#1F4E79")
        normal_style = styles["Normal"]

        # --- Création PDF ---
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
        doc = SimpleDocTemplate(pdf_path, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=3*cm, bottomMargin=2*cm)
        elements = []
        title_keywords = [
            'résumé exécutif', 'méthodologie', 'section social', 'section environnement','performance rse','conclusion',
            'gouvernance', 'alignement avec les objectifs de développement durable (odd)' ,'performance de production et consommation','recommandations'
        ]

        for sec in content.split("\n\n"):
            sec = sec.strip()
            if not sec:
                continue

            table_data = parse_table(sec)
            if table_data:
                # Style du tableau : header si présent
                t = Table(table_data, hAlign='LEFT')
                # si on a au moins 2 lignes, traiter la première comme en-tête
                has_header = len(table_data) >= 1
                ts = TableStyle([
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('LEFTPADDING', (0,0), (-1,-1), 6),
                    ('RIGHTPADDING', (0,0), (-1,-1), 6),
                ])
                if has_header:
                    ts.add('BACKGROUND', (0, 0), (-1, 0), colors.lightblue)
                    ts.add('TEXTCOLOR', (0, 0), (-1, 0), colors.black)
                    ts.add('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold')
                t.setStyle(ts)
                elements.append(t)
                elements.append(Spacer(1, 12))
                continue

            # titres
            if any(sec.lower().startswith(k) for k in title_keywords):
                elements.append(Paragraph(sec, title_style))
            else:
                elements.append(Paragraph(sec, normal_style))
            elements.append(Spacer(1, 12))

        # build PDF (callbacks pour pagination)
        doc.build(elements, onFirstPage=add_page_number, onLaterPages=add_page_number)

        # --- Création Word (.docx) ---
        word_path = os.path.join(output_dir, f"{base_name}.docx")
        document = Document()
        for block in content.split("\n\n"):
            line_strip = block.strip()
            if not line_strip:
                continue
            table_data = parse_table(line_strip)
            if table_data:
                table = document.add_table(rows=len(table_data), cols=len(table_data[0]))
                for i, row in enumerate(table_data):
                    for j, cell_text in enumerate(row):
                        cell = table.cell(i, j)
                        cell.text = cell_text
                        # mettre le header en gras
                        if i == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                # ajouter un paragraphe vide après le tableau
                document.add_paragraph()
                continue

            if any(line_strip.lower().startswith(k) for k in title_keywords):
                document.add_heading(line_strip, level=2)
            else:
                document.add_paragraph(line_strip, style="Normal")

        document.save(word_path)

        # --- Création TXT ---
        txt_path = os.path.join(output_dir, f"{base_name}.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(content)

        # --- URLs ---
        pdf_url = f"{settings.MEDIA_URL}reports/{base_name}.pdf"
        word_url = f"{settings.MEDIA_URL}reports/{base_name}.docx"
        txt_url = f"{settings.MEDIA_URL}reports/{base_name}.txt"

        return JsonResponse({
            "pdf_url": pdf_url,
            "word_url": word_url,
            "txt_url": txt_url
        })

    return render(request, "reports/generate_report.html")